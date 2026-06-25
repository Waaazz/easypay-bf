from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Color palette ─────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1e, 0x3a, 0x8a)
BLUE_MID   = RGBColor(0x2d, 0x6a, 0xd6)
BLUE_LIGHT = RGBColor(0xdb, 0xe9, 0xfe)
GREEN      = RGBColor(0x16, 0xa3, 0x4a)
ORANGE     = RGBColor(0xf9, 0x73, 0x16)
PURPLE     = RGBColor(0x7c, 0x3a, 0xed)
GRAY_DARK  = RGBColor(0x1f, 0x29, 0x37)
GRAY_MED   = RGBColor(0x6b, 0x72, 0x80)
WHITE      = RGBColor(0xff, 0xff, 0xff)
RED        = RGBColor(0xdc, 0x26, 0x26)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_run_font(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color

def heading(text, level=1, color=BLUE_DARK, size=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(size or sizes.get(level, 12))
    run.font.color.rgb = color
    run.font.name = "Calibri"
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, indent=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0, bold_part=None, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    if bold_part and text.startswith(bold_part):
        r1 = p.add_run(bold_part)
        r1.font.bold = True
        r1.font.size = Pt(11)
        r1.font.name = "Calibri"
        r1.font.color.rgb = BLUE_DARK
        rest = text[len(bold_part):]
        r2 = p.add_run(rest)
        r2.font.size = Pt(11)
        r2.font.name = "Calibri"
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = color
    return p

def shade_cell(cell, color_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=WHITE, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.bold  = bold
    run.font.size  = Pt(size)
    run.font.name  = "Calibri"
    run.font.color.rgb = color

def add_info_table(rows_data):
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        shade_cell(row.cells[0], "1e3a8a")
        set_cell_text(row.cells[0], label, bold=True, color=WHITE, size=10)
        row.cells[0].width = Cm(5)
        row.cells[1].text = value
        for para in row.cells[1].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1e3a8a")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)

def make_table(headers, rows, header_color="1e3a8a", alt_color="f8faff"):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        shade_cell(t.rows[0].cells[i], header_color)
        set_cell_text(t.rows[0].cells[i], h, bold=True, color=WHITE, size=10)
    for ri, row_data in enumerate(rows, 1):
        if ri % 2 == 0:
            for ci in range(len(headers)):
                shade_cell(t.rows[ri].cells[ci], alt_color)
        for ci, val in enumerate(row_data):
            t.rows[ri].cells[ci].text = val
            for para in t.rows[ri].cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = "Calibri"
                    run.font.color.rgb = GRAY_DARK
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════════
#  PAGE DE COUVERTURE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph("\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EASYPAY BF")
run.font.size  = Pt(36)
run.font.bold  = True
run.font.name  = "Calibri"
run.font.color.rgb = BLUE_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Plateforme de gestion de dépôts et retraits")
run.font.size  = Pt(18)
run.font.name  = "Calibri"
run.font.color.rgb = BLUE_MID

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("─────────────────────────────────────────")
run.font.color.rgb = BLUE_MID
run.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CAHIER DES CHARGES")
run.font.size  = Pt(28)
run.font.bold  = True
run.font.name  = "Calibri"
run.font.color.rgb = GRAY_DARK

doc.add_paragraph("\n\n\n")

add_info_table([
    ("Version",    "2.0"),
    ("Date",       datetime.date.today().strftime("%d %B %Y")),
    ("Statut",     "En cours — implémentation avancée"),
    ("Projet",     "EasyPay BF"),
    ("Domaine",    "Paris sportifs / Mobile Money — Burkina Faso"),
    ("Stack tech", "React 18 · Firebase · Tailwind CSS · Vite · Vercel"),
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE DES MATIÈRES
# ═══════════════════════════════════════════════════════════════════════════════
heading("TABLE DES MATIÈRES", 1)
toc = [
    "1.  Contexte et présentation du projet",
    "2.  Objectifs",
    "3.  Parties prenantes",
    "4.  Périmètre fonctionnel",
    "    4.1  Page d'accueil (Client)",
    "    4.2  Module Dépôt — flux en 6 étapes",
    "    4.3  Module Retrait — flux en 6 étapes",
    "    4.4  Historique des transactions (Client)",
    "    4.5  Espace Agent",
    "    4.6  Espace Admin",
    "5.  Authentification et gestion des accès",
    "6.  Règles métier",
    "7.  Architecture technique",
    "8.  Modèle de données",
    "9.  Routes et interfaces",
    "10. Sécurité et conformité",
    "11. Contraintes non-fonctionnelles",
    "12. Livrables",
    "13. Points ouverts",
]
for line in toc:
    p = doc.add_paragraph()
    run = p.add_run(line)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    if line.startswith("    "):
        p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. CONTEXTE
# ═══════════════════════════════════════════════════════════════════════════════
heading("1. Contexte et présentation du projet", 1)
divider()
body(
    "EasyPay BF est une application web mobile-first déployée au Burkina Faso. "
    "Elle sert d'intermédiaire entre les joueurs de plateformes de paris sportifs "
    "(1XBET, MELBET, BETWINNER) et des agents locaux qui traitent les opérations "
    "financières via les solutions de paiement mobile (Orange Money, Telmob, Telecel)."
)
body(
    "Les joueurs ne disposent souvent pas de moyens pratiques pour recharger leurs "
    "comptes de jeu ou en retirer les gains directement. EasyPay BF centralise ces "
    "demandes et les route vers des agents terrain qui exécutent les transferts."
)
body(
    "Les agents sont recrutés et gérés par l'administrateur de la plateforme. "
    "Chaque agent dispose d'un compte dédié (numéro + mot de passe) créé par l'admin."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  2. OBJECTIFS
# ═══════════════════════════════════════════════════════════════════════════════
heading("2. Objectifs", 1)
divider()
bullet("Permettre aux clients de soumettre des demandes de dépôt et retrait 24h/24, 7j/7.")
bullet("Fournir aux agents un espace de travail pour accepter, traiter et clôturer les commandes.")
bullet("Permettre à l'administrateur de créer et gérer les comptes agents.")
bullet("Offrir à l'administrateur une vue analytique complète de l'activité.")
bullet("Réduire les erreurs et frictions liées aux transactions manuelles (WhatsApp, appels).")
bullet("Assurer la traçabilité complète de chaque transaction.")

# ═══════════════════════════════════════════════════════════════════════════════
#  3. PARTIES PRENANTES
# ═══════════════════════════════════════════════════════════════════════════════
heading("3. Parties prenantes", 1)
divider()

make_table(
    ["Rôle", "Description", "Mode d'accès"],
    [
        ("Client",
         "Joueur souhaitant déposer ou retirer de l'argent. Aucune inscription requise.",
         "Navigation anonyme (Firebase Anonymous Auth). Historique lié à l'appareil."),
        ("Agent",
         "Opérateur terrain qui exécute les transferts Mobile Money.",
         "Connexion par numéro de téléphone + mot de passe sur /agent/login"),
        ("Admin",
         "Gestionnaire de la plateforme. Crée les agents, supervise l'activité.",
         "Accès direct /admin (sécurisation à renforcer)"),
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  4. PÉRIMÈTRE FONCTIONNEL
# ═══════════════════════════════════════════════════════════════════════════════
heading("4. Périmètre fonctionnel", 1)
divider()

# 4.1
heading("4.1 Page d'accueil (Client)", 2)
body("Point d'entrée principal de l'application. Le client :")
bullet("Choisit la plateforme de jeu : 1XBET, MELBET ou BETWINNER.")
bullet("Saisit le montant via un pavé numérique (minimum 500 FCFA, maximum 9 999 999 FCFA).")
bullet("Lance une opération Dépôt ou Retrait.")
bullet("Peut accéder au support WhatsApp.")

# 4.2
heading("4.2 Module Dépôt — flux en 6 étapes", 2)
body("Flux complet de dépôt implémenté, inspiré de l'UX Futurix Pay :")

make_table(
    ["Étape", "Écran", "Description"],
    [
        ("1", "ID du compte",
         "Le client saisit son identifiant de compte sur la plateforme choisie. "
         "Avertissement : aucun remboursement en cas d'erreur d'ID."),
        ("2", "Choix opérateur",
         "Le client choisit l'opérateur Mobile Money : Orange Money, Telmob ou Telecel. "
         "Chaque opérateur affiche son numéro agent dédié."),
        ("3", "Numéro de téléphone",
         "Le client saisit le numéro de téléphone depuis lequel il va envoyer le paiement."),
        ("4", "Paiement",
         "Affichage du code USSD à composer et du numéro agent. "
         "Orange Money : *144*2*1*{agent}*{montant}# · "
         "Telmob : *555*1*{agent}*{montant}# · "
         "Telecel : *135*1*{agent}*{montant}#. "
         "Bouton « J'ai envoyé le paiement » → statut passe à awaiting_confirmation."),
        ("5", "Attente",
         "Minuterie de session (15 min). Suivi en temps réel du statut via onSnapshot. "
         "L'écran se met à jour automatiquement dès que l'agent confirme."),
        ("6", "Succès / Annulé",
         "Confirmation visuelle selon le résultat de l'agent."),
    ]
)

body("Note : la transaction est créée dans Firestore à l'étape 3 (saisie du numéro de téléphone) "
     "avec le statut « pending ».", color=BLUE_MID)

# 4.3
heading("4.3 Module Retrait — flux en 6 étapes", 2)
body("Flux complet de retrait avec aperçu des frais :")

make_table(
    ["Étape", "Écran", "Description"],
    [
        ("1", "ID du compte",
         "Le client saisit son identifiant de compte sur la plateforme choisie."),
        ("2", "Infos Orange Money",
         "Seul Orange Money est accepté pour les retraits. "
         "Le client saisit son numéro et le nom du titulaire du compte."),
        ("3", "Code de retrait",
         "Le client génère un code de retrait sur la plateforme de jeu "
         "(menu « Express Multi service ») et le saisit ici."),
        ("4", "Récapitulatif + Frais",
         "Affichage du récapitulatif complet : montant demandé, frais (1%), "
         "montant net que le client recevra. Le client confirme la demande."),
        ("5", "Attente",
         "Suivi en temps réel du statut via onSnapshot."),
        ("6", "Succès / Annulé",
         "Confirmation visuelle selon le résultat."),
    ]
)

body("Note : la transaction est créée à l'étape 4 (confirmation du récapitulatif) "
     "avec le statut « pending ». Frais de retrait : 1 % du montant.", color=BLUE_MID)

# 4.4
heading("4.4 Historique des transactions (Client)", 2)
body("Accessible depuis le menu de navigation (/history) :")
bullet("Affiche les 50 dernières transactions liées à l'UID Firebase anonyme du client.")
bullet("Filtres par type (dépôt / retrait) et par statut.")
bullet("Affiche : montant, référence, date, statut, opérateur, plateforme.")
bullet("Fonctionne grâce à l'authentification anonyme persistante (même navigateur).")

# 4.5
heading("4.5 Espace Agent", 2)
body("L'agent accède à son espace depuis /agent/login avec son numéro + mot de passe.")

heading("4.5.1 Tableau de bord (/agent)", 3)
bullet("Vue temps réel de toutes les commandes actives (pending + awaiting_confirmation + processing).")
bullet("Compteurs : commandes en attente / en cours.")
bullet("Onglets : Tout / Attente / En cours.")
bullet("Alerte violette quand le client a confirmé l'envoi du paiement (awaiting_confirmation).")

heading("Actions sur une commande :", 3)
bullet("Accepter / Vérifier et Accepter → statut processing.", level=1)
bullet("Marquer terminé → statut completed.", level=1)
bullet("Annuler → statut cancelled + raison obligatoire.", level=1)

heading("4.5.2 Historique commandes (/agent/orders)", 3)
bullet("Affiche uniquement les commandes traitées par cet agent (filtrées par agentId).")
bullet("Onglets : Terminées / Annulées.")
bullet("Statistiques : nombre terminées, annulées, volume total traité.")

# 4.6
heading("4.6 Espace Admin", 2)
body("Menu de navigation en sidebar avec 3 sections :")

make_table(
    ["Section", "Page", "Contenu"],
    [
        ("Général", "/admin",
         "6 KPIs : total transactions, volume dépôts, volume retraits, complétées, agents actifs, volume total. "
         "Graphique en barres 7 jours (dépôts vs retraits). 5 dernières transactions."),
        ("Activité", "/admin/transactions",
         "Liste complète avec recherche (client/téléphone/ID), filtres statut et type."),
        ("Activité", "/admin/deposits",
         "Même vue, pré-filtrée sur les dépôts uniquement."),
        ("Activité", "/admin/withdrawals",
         "Même vue, pré-filtrée sur les retraits uniquement."),
        ("Équipe", "/admin/agents",
         "Liste des agents (actifs/inactifs). Création de compte agent par l'admin. Activation/désactivation."),
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  5. AUTHENTIFICATION ET GESTION DES ACCÈS
# ═══════════════════════════════════════════════════════════════════════════════
heading("5. Authentification et gestion des accès", 1)
divider()

heading("5.1 Clients (Firebase Anonymous Auth)", 2)
body(
    "Les clients ne s'inscrivent pas. À la première visite, Firebase crée automatiquement "
    "un compte anonyme persistant lié au navigateur. Cet UID permet d'afficher "
    "l'historique des transactions du client lors de ses prochaines visites."
)
bullet("Provider Firebase : Anonymous Authentication")
bullet("Avantage : aucune friction pour le client, historique persistant.")
bullet("Limite : l'historique est perdu si l'utilisateur efface ses données de navigation.")

heading("5.2 Agents (Email + Mot de passe)", 2)
body("Les agents ne créent pas eux-mêmes leur compte. Le flux est le suivant :")
bullet("L'admin va sur /admin/agents → « Créer un agent ».", bold_part="Étape 1 — Admin : ")
bullet("Il saisit le nom complet, le numéro de téléphone et un mot de passe.", bold_part="Étape 2 — Admin : ")
bullet(
    "Le système crée un compte Firebase Email/Password avec l'email synthétique "
    "226{numéro}@easypay-agent.bf via une instance Firebase secondaire "
    "(sans déconnecter l'admin).", bold_part="Étape 3 — Système : "
)
bullet("L'admin communique le numéro + mot de passe à l'agent.", bold_part="Étape 4 — Admin : ")
bullet("L'agent se connecte sur /agent/login avec son numéro + mot de passe.", bold_part="Étape 5 — Agent : ")

body("⚠️ Prérequis Firebase : activer le provider Email/Password dans Authentication → Sign-in method.", color=ORANGE)

heading("5.3 Admin", 2)
body(
    "L'accès admin est actuellement non sécurisé (accès direct à /admin). "
    "Une sécurisation par email/password + vérification du rôle est prévue."
)
bullet("Action requise : créer un compte admin dans Firebase et ajouter un guard sur /admin.", color=ORANGE)

# ═══════════════════════════════════════════════════════════════════════════════
#  6. RÈGLES MÉTIER
# ═══════════════════════════════════════════════════════════════════════════════
heading("6. Règles métier", 1)
divider()

make_table(
    ["#", "Règle", "Valeur / Condition"],
    [
        ("R1",  "Montant minimum",                  "500 FCFA"),
        ("R2",  "Montant maximum",                  "9 999 999 FCFA (7 chiffres)"),
        ("R3",  "Frais de dépôt",                   "0 % (gratuit pour le client)"),
        ("R4",  "Frais de retrait",                 "1 % du montant (configurable dans constants.js → FEES.withdrawal)"),
        ("R5",  "Opérateurs dépôt",                 "Orange Money, Telmob, Telecel"),
        ("R6",  "Opérateurs retrait",               "Orange Money UNIQUEMENT"),
        ("R7",  "Plateformes supportées",            "1XBET, MELBET, BETWINNER"),
        ("R8",  "Durée de session dépôt",           "15 minutes (minuterie affichée au client)"),
        ("R9",  "Code USSD Orange Money",           "*144*2*1*{numéroAgent}*{montant}#"),
        ("R10", "Code USSD Telmob",                 "*555*1*{numéroAgent}*{montant}#"),
        ("R11", "Code USSD Telecel",                "*135*1*{numéroAgent}*{montant}#"),
        ("R12", "Historique client",                "50 dernières transactions (limit Firestore)"),
        ("R13", "Historique admin",                 "100 dernières transactions (limit Firestore)"),
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  7. ARCHITECTURE TECHNIQUE
# ═══════════════════════════════════════════════════════════════════════════════
heading("7. Architecture technique", 1)
divider()

heading("Frontend", 2)
bullet("Framework : React 18 avec Vite (bundler)")
bullet("Routing : React Router v6")
bullet("UI : Tailwind CSS v3 (mobile-first, dark mode pour agent/admin)")
bullet("Graphiques : Recharts (BarChart temps réel)")
bullet("Icônes : Lucide React")

heading("Backend / Infrastructure", 2)
bullet("Base de données : Firebase Firestore (NoSQL, temps réel via onSnapshot)")
bullet("Auth clients : Firebase Anonymous Authentication")
bullet("Auth agents : Firebase Email/Password (email synthétique 226xxxx@easypay-agent.bf)")
bullet("Instance Firebase secondaire pour créer des agents sans déconnecter l'admin")
bullet("Hébergement : Vercel (SPA, routes gérées via vercel.json)")
bullet("Variables d'environnement : VITE_FIREBASE_* (non committées, fichier .env)")

heading("Flux de données", 2)
body(
    "Client → soumet formulaire → createTransaction() → Firestore (pending) → "
    "Agent reçoit en temps réel (onSnapshot) → acceptOrder() [processing] → "
    "completeOrder() [completed] | cancelOrder() [cancelled]."
)
body(
    "Pour les dépôts : le client envoie le paiement USSD, clique « J'ai envoyé » → "
    "confirmPayment() [awaiting_confirmation] → l'agent voit l'alerte violette et vérifie."
)

heading("Tri et requêtes Firestore", 2)
body(
    "Les requêtes Firestore n'utilisent pas orderBy (évite les index composites). "
    "Le tri par date est effectué côté client (.sort() sur createdAt.seconds)."
)

# ═══════════════════════════════════════════════════════════════════════════════
#  8. MODÈLE DE DONNÉES
# ═══════════════════════════════════════════════════════════════════════════════
heading("8. Modèle de données", 1)
divider()

heading("8.1 Collection : transactions", 2)

make_table(
    ["Champ", "Type", "Présent si", "Description"],
    [
        ("id",           "string (auto)",  "Toujours",       "Identifiant Firestore auto-généré"),
        ("type",         "string",         "Toujours",       "« deposit » ou « withdrawal »"),
        ("amount",       "number",         "Toujours",       "Montant en FCFA"),
        ("platform",     "string",         "Toujours",       "« 1xbet », « melbet » ou « betwinner »"),
        ("accountId",    "string",         "Toujours",       "Identifiant du compte de jeu"),
        ("operator",     "string",         "Toujours",       "« orange », « telmob » ou « telecel »"),
        ("clientPhone",  "string",         "Dépôt",          "Numéro depuis lequel le paiement est envoyé"),
        ("phone",        "string",         "Retrait",        "Numéro Orange Money de réception"),
        ("accountName",  "string",         "Retrait",        "Nom du titulaire Orange Money"),
        ("withdrawCode", "string",         "Retrait",        "Code de retrait Express Multi service"),
        ("status",       "string",         "Toujours",       "Voir statuts ci-dessous"),
        ("clientId",     "string",         "Toujours",       "UID Firebase du client (anonymous ou réel)"),
        ("clientName",   "string",         "Toujours",       "Nom ou téléphone du client"),
        ("agentId",      "string | null",  "Après acceptation","UID Firebase de l'agent"),
        ("note",         "string | null",  "Annulation",     "Raison de l'annulation"),
        ("createdAt",    "Timestamp",      "Toujours",       "Date/heure de création (serverTimestamp)"),
        ("updatedAt",    "Timestamp",      "Toujours",       "Date/heure de dernière modification"),
    ]
)

heading("8.2 Statuts des transactions", 2)

make_table(
    ["Statut", "Libellé affiché", "Couleur", "Signification"],
    [
        ("pending",                "En attente",      "Jaune",  "Transaction créée, aucun agent n'a encore pris en charge"),
        ("awaiting_confirmation",  "Paiement envoyé", "Violet", "Le client a cliqué « J'ai envoyé le paiement » — l'agent doit vérifier son compte Mobile Money"),
        ("processing",             "En cours",        "Bleu",   "Agent a accepté la commande, traitement en cours"),
        ("completed",              "Terminé",         "Vert",   "Transaction confirmée par l'agent"),
        ("cancelled",              "Annulé",          "Rouge",  "Annulation avec raison obligatoire"),
    ]
)

heading("8.3 Collection : users", 2)

make_table(
    ["Champ", "Type", "Description"],
    [
        ("uid",       "string",    "UID Firebase Authentication"),
        ("name",      "string",    "Nom complet de l'agent"),
        ("phone",     "string",    "Numéro au format +226XXXXXXXX"),
        ("email",     "string",    "Email synthétique : 226XXXXXXXX@easypay-agent.bf"),
        ("role",      "string",    "« client », « agent » ou « admin »"),
        ("active",    "boolean",   "Compte actif ou désactivé par l'admin"),
        ("createdAt", "Timestamp", "Date de création du compte"),
        ("updatedAt", "Timestamp", "Dernière modification"),
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  9. ROUTES ET INTERFACES
# ═══════════════════════════════════════════════════════════════════════════════
heading("9. Routes et interfaces", 1)
divider()

make_table(
    ["Route", "Accès", "Description"],
    [
        ("/",                    "Client",  "Page d'accueil : choix plateforme + saisie montant"),
        ("/deposit",             "Client",  "Flux dépôt 6 étapes"),
        ("/withdrawal",          "Client",  "Flux retrait 6 étapes"),
        ("/history",             "Client",  "Historique des 50 dernières transactions du client"),
        ("/login",               "Client",  "Connexion client par numéro OTP (optionnel)"),
        ("/agent/login",         "Agent",   "Connexion agent : numéro + mot de passe"),
        ("/agent",               "Agent",   "Dashboard agent : commandes actives temps réel"),
        ("/agent/orders",        "Agent",   "Historique des commandes traitées par cet agent"),
        ("/agent/order/:id",     "Agent",   "Détail d'une commande"),
        ("/admin",               "Admin",   "Dashboard : KPIs + graphique + transactions récentes"),
        ("/admin/transactions",  "Admin",   "Liste complète toutes transactions"),
        ("/admin/deposits",      "Admin",   "Liste filtrée sur les dépôts"),
        ("/admin/withdrawals",   "Admin",   "Liste filtrée sur les retraits"),
        ("/admin/agents",        "Admin",   "Gestion agents : créer, activer, désactiver"),
    ]
)

heading("Numéros agent par opérateur", 2)
body("Configurés dans src/utils/constants.js → AGENT_NUMBERS (à mettre à jour avec les vrais numéros) :")

make_table(
    ["Opérateur", "Numéro agent (placeholder)", "Code USSD"],
    [
        ("Orange Money", "+226 07 00 00 00", "*144*2*1*07000000*{montant}#"),
        ("Telmob",       "+226 60 00 00 00", "*555*1*60000000*{montant}#"),
        ("Telecel",      "+226 55 00 00 00", "*135*1*55000000*{montant}#"),
    ]
)
body("⚠️ Ces numéros sont des placeholders. Remplacer par les vrais numéros agents dans constants.js.", color=ORANGE)

# ═══════════════════════════════════════════════════════════════════════════════
#  10. SÉCURITÉ ET CONFORMITÉ
# ═══════════════════════════════════════════════════════════════════════════════
heading("10. Sécurité et conformité", 1)
divider()

bullet("Authentification Firebase (JWT) — aucun mot de passe en clair dans Firestore.")
bullet("Instance Firebase secondaire pour créer des agents : l'admin reste connecté pendant l'opération.")
bullet("Variables d'environnement Firebase stockées dans .env (non versionnées dans Git).")
bullet("Tri des données côté client (pas d'index composites Firestore requis).")

body("⚠️  Points à implémenter :", color=ORANGE)
bullet("Firestore Security Rules : restreindre lecture/écriture par rôle.", level=1, color=ORANGE)
bullet("Route guards (ProtectedRoute) sur /agent et /admin.", level=1, color=ORANGE)
bullet("Sécuriser l'accès admin par authentification dédiée.", level=1, color=ORANGE)
bullet("Validation côté serveur des montants et transitions de statut (Cloud Functions).", level=1, color=ORANGE)

# ═══════════════════════════════════════════════════════════════════════════════
#  11. CONTRAINTES NON-FONCTIONNELLES
# ═══════════════════════════════════════════════════════════════════════════════
heading("11. Contraintes non-fonctionnelles", 1)
divider()

make_table(
    ["Exigence", "Cible", "Commentaire"],
    [
        ("Performance",    "< 3 s chargement initial",      "Vite + Vercel CDN"),
        ("Temps réel",     "< 2 s délai de mise à jour",    "Firestore onSnapshot"),
        ("Disponibilité",  "≥ 99 % (SLA Vercel/Firebase)",  "Monitoring à prévoir"),
        ("Compatibilité",  "Chrome, Firefox, Safari mobile", "iOS 14+ / Android 9+"),
        ("Accessibilité",  "Contrastes WCAG AA",             "Dark mode agent/admin"),
        ("Scalabilité",    "Firestore auto-scale",           "Limite 100 tx/lecture admin"),
        ("Langue",         "Français (fr-FR)",               "Formatage monétaire FCFA"),
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  12. LIVRABLES
# ═══════════════════════════════════════════════════════════════════════════════
heading("12. Livrables", 1)
divider()
bullet("Application web déployée sur Vercel (URL de production).")
bullet("Code source versionné sur Git (branche master).")
bullet("Configuration Firebase (Firestore rules, Auth providers activés).")
bullet("Comptes agents créés via l'interface admin.")
bullet("Ce cahier des charges v2.0 validé.")

# ═══════════════════════════════════════════════════════════════════════════════
#  13. POINTS OUVERTS
# ═══════════════════════════════════════════════════════════════════════════════
heading("13. Points ouverts / À définir", 1)
divider()

make_table(
    ["#", "Point ouvert", "Questions / Actions"],
    [
        ("1",  "Vrais numéros agents",
               "Remplacer les placeholders dans constants.js → AGENT_NUMBERS (Orange, Telmob, Telecel)."),
        ("2",  "Numéro WhatsApp support",
               "Remplacer WHATSAPP_NUMBER = '22600000000' dans constants.js."),
        ("3",  "Codes USSD Telmob/Telecel",
               "Vérifier les formats USSD exacts : *555*... et *135*... avec les opérateurs."),
        ("4",  "Sécurisation /admin",
               "Ajouter un login admin dédié et un route guard."),
        ("5",  "Firestore Security Rules",
               "Définir les règles par rôle pour chaque collection."),
        ("6",  "Route guards /agent",
               "Rediriger vers /agent/login si l'utilisateur n'est pas authentifié comme agent."),
        ("7",  "Minimum dépôt",
               "Confirmer : 500 FCFA (actuellement dans le pavé numérique)."),
        ("8",  "Nouvelles plateformes",
               "Ajouter d'autres plateformes ? (ex : Bet9ja, SportyBet)"),
        ("9",  "Notifications push/SMS",
               "Notifier le client quand sa transaction est traitée ?"),
        ("10", "Export CSV/PDF",
               "L'admin a-t-il besoin d'exporter les transactions ?"),
        ("11", "Programme de parrainage",
               "Bouton visible sur l'accueil. Inclure dans le périmètre ou supprimer ?"),
        ("12", "CGU et confidentialité",
               "Documents légaux à rédiger et lier dans l'application."),
    ]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  PAGE DE SIGNATURE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("Validation du document", 1, align=WD_ALIGN_PARAGRAPH.CENTER)
divider()

doc.add_paragraph()
sig = doc.add_table(rows=3, cols=2)
sig.style = "Table Grid"
for i, h in enumerate(["Rédacteur / Porteur du projet", "Validateur"]):
    shade_cell(sig.rows[0].cells[i], "1e3a8a")
    set_cell_text(sig.rows[0].cells[i], h, bold=True, color=WHITE, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

sig.rows[1].cells[0].text = "Nom : ___________________________"
sig.rows[1].cells[1].text = "Nom : ___________________________"
sig.rows[2].cells[0].text = "Date : __________________________\n\nSignature : "
sig.rows[2].cells[1].text = "Date : __________________________\n\nSignature : "
for row in sig.rows[1:]:
    for cell in row.cells:
        cell.height = Cm(3)
        for para in cell.paragraphs:
            para.paragraph_format.space_before = Pt(6)
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = "Calibri"

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EasyPay BF — Cahier des charges v2.0 — " + datetime.date.today().strftime("%d/%m/%Y"))
run.font.size = Pt(9)
run.font.color.rgb = GRAY_MED
run.font.name = "Calibri"

# ── Save ──────────────────────────────────────────────────────────────────────
out = r"C:\Users\Z book\Desktop\EasyPay_BF_Cahier_des_charges.docx"
doc.save(out)
print(f"Document genere : {out}")
