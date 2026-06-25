from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Marges ────────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Palette ───────────────────────────────────────────────────────────────────
BLUE      = RGBColor(0x25, 0x63, 0xeb)
BLUE_DARK = RGBColor(0x1e, 0x3a, 0x8a)
GREEN     = RGBColor(0x16, 0xa3, 0x4a)
RED       = RGBColor(0xdc, 0x26, 0x26)
GRAY      = RGBColor(0x6b, 0x72, 0x80)
DARK      = RGBColor(0x1a, 0x1a, 0x1a)
WHITE     = RGBColor(0xff, 0xff, 0xff)
YELLOW_BG = "fffbeb"
GREEN_BG  = "f0fdf4"
BLUE_BG   = "f8faff"

TODAY = datetime.date.today().strftime("%d %B %Y")
TODAY_SHORT = datetime.date.today().strftime("%d/%m/%Y")

# ── Helpers ───────────────────────────────────────────────────────────────────
def shade_cell(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, color=DARK, size=10,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.size   = Pt(size)
    run.font.name   = "Calibri"
    run.font.color.rgb = color

def cell_multiline(cell, lines, size=10):
    """lines: list of (text, bold, color)"""
    cell.text = ""
    for i, (text, bold, color) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(text)
        run.font.bold  = bold
        run.font.size  = Pt(size)
        run.font.name  = "Calibri"
        run.font.color.rgb = color

def section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.bold  = True
    run.font.size  = Pt(12)
    run.font.name  = "Calibri"
    run.font.color.rgb = BLUE
    # Ligne sous le titre
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "e5e7eb")
    pBdr.append(bot)
    pPr.append(pBdr)

def small_text(text, color=GRAY, size=10, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.name   = "Calibri"
    run.font.italic = italic
    run.font.color.rgb = color

def add_table(headers, rows, col_widths=None,
              header_bg="1e3a8a", alt_bg="f9fafb",
              footer_rows=None):
    t = doc.add_table(rows=1 + len(rows) + (len(footer_rows) if footer_rows else 0),
                      cols=len(headers))
    t.style = "Table Grid"

    # Header
    for i, h in enumerate(headers):
        shade_cell(t.rows[0].cells[i], header_bg)
        cell_text(t.rows[0].cells[i], h, bold=True, color=WHITE, size=10)

    # Body
    for ri, row in enumerate(rows, 1):
        if ri % 2 == 0:
            for ci in range(len(headers)):
                shade_cell(t.rows[ri].cells[ci], alt_bg)
        if isinstance(row[0], list):
            # Multi-line first cell
            for ci, cell_data in enumerate(row):
                if isinstance(cell_data, list):
                    cell_multiline(t.rows[ri].cells[ci], cell_data)
                else:
                    cell_text(t.rows[ri].cells[ci], cell_data)
        else:
            for ci, val in enumerate(row):
                align = WD_ALIGN_PARAGRAPH.RIGHT if ci == len(row)-1 and '000' in str(val) else WD_ALIGN_PARAGRAPH.LEFT
                bold = ci == len(row)-1 and 'FCFA' in str(val)
                color = BLUE if bold else DARK
                cell_text(t.rows[ri].cells[ci], str(val), bold=bold, color=color,
                          align=align, size=10)

    # Footer rows
    if footer_rows:
        base = 1 + len(rows)
        for fi, fr in enumerate(footer_rows):
            for ci in range(len(headers)):
                shade_cell(t.rows[base+fi].cells[ci], "1e3a8a")
            for ci, val in enumerate(fr):
                align = WD_ALIGN_PARAGRAPH.RIGHT if ci == len(fr)-1 else WD_ALIGN_PARAGRAPH.LEFT
                cell_text(t.rows[base+fi].cells[ci], val, bold=True,
                          color=WHITE, size=10, align=align)

    if col_widths:
        for ri2 in range(len(t.rows)):
            for ci2, w in enumerate(col_widths):
                t.rows[ri2].cells[ci2].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

# ═══════════════════════════════════════════════════════════════════════════════
#  EN-TÊTE DU DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════

# Barre meta
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
run = p.add_run(
    f"DEVIS  |  Application Web EasyPay BF  |  "
    f"Réf. WAA-EASYPAY-2026-001  |  {TODAY}  |  Document confidentiel"
)
run.font.size   = Pt(8)
run.font.name   = "Calibri"
run.font.italic = True
run.font.color.rgb = GRAY

# Tableau en-tête  (prestataire | DEVIS)
ht = doc.add_table(rows=1, cols=2)
ht.style = "Table Grid"
left  = ht.rows[0].cells[0]
right = ht.rows[0].cells[1]

# Cellule gauche — infos prestataire
left_p = left.paragraphs[0]
def add_left(text, bold=False, size=11, color=DARK):
    run = left_p.add_run(('\n' if left_p.runs else '') + text)
    run.font.bold  = bold
    run.font.size  = Pt(size)
    run.font.name  = "Calibri"
    run.font.color.rgb = color

left.paragraphs[0].paragraph_format.space_before = Pt(4)
r = left.paragraphs[0].add_run("WOUBI ABDOUL AZIZ")
r.font.bold = True; r.font.size = Pt(14); r.font.name = "Calibri"; r.font.color.rgb = DARK

for line_txt, b, sz, col in [
    ("Développeur Full-Stack & Mobile", False, 10, GRAY),
    ("", False, 6, DARK),
    ("woubiaziz@gmail.com", False, 10, DARK),
    ("+226 54 12 56 37 / +226 68 87 71 82", False, 10, DARK),
    ("Ouagadougou, Burkina Faso", False, 10, DARK),
    ("Licence en Génie Logiciel — UVBF", False, 10, DARK),
]:
    p2 = left.add_paragraph(line_txt)
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after  = Pt(1)
    for run2 in p2.runs:
        run2.font.bold  = b
        run2.font.size  = Pt(sz)
        run2.font.name  = "Calibri"
        run2.font.color.rgb = col

# Cellule droite — DEVIS + ref
rp = right.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r2 = rp.add_run("DEVIS")
r2.font.bold = True; r2.font.size = Pt(32); r2.font.name = "Calibri"; r2.font.color.rgb = DARK

for line_txt, b, sz, col in [
    ("Réf. : WAA-EASYPAY-2026-001", True, 10, BLUE),
    (f"Date : {TODAY}", False, 10, GRAY),
    ("Validité : 30 jours", False, 10, GRAY),
    ("", False, 6, DARK),
    ("CLIENT", True, 8, GRAY),
    ("[NOM DU CLIENT]", True, 12, DARK),
    ("Ouagadougou, Burkina Faso", False, 10, GRAY),
]:
    p3 = right.add_paragraph(line_txt)
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.paragraph_format.space_before = Pt(1)
    p3.paragraph_format.space_after  = Pt(1)
    for run3 in p3.runs:
        run3.font.bold  = b
        run3.font.size  = Pt(sz)
        run3.font.name  = "Calibri"
        run3.font.color.rgb = col

ht.rows[0].cells[0].width = Cm(9)
ht.rows[0].cells[1].width = Cm(7)
shade_cell(left, "ffffff")
shade_cell(right, "f8faff")

# Séparateur bleu
sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(8)
sep.paragraph_format.space_after  = Pt(8)
pPr = sep._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot = OxmlElement("w:bottom")
bot.set(qn("w:val"),   "single")
bot.set(qn("w:sz"),    "12")
bot.set(qn("w:space"), "1")
bot.set(qn("w:color"), "2563eb")
pBdr.append(bot)
pPr.append(pBdr)

# ── OBJET ─────────────────────────────────────────────────────────────────────
section_title("Objet de la prestation")

ob = doc.add_table(rows=1, cols=1)
ob.style = "Table Grid"
shade_cell(ob.rows[0].cells[0], BLUE_BG.replace("#",""))
c = ob.rows[0].cells[0]

lines_ob = [
    ("Développement d'une application web mobile-first (PWA) pour EasyPay BF, "
     "une plateforme de dépôt et de retrait d'argent dédiée aux paris sportifs au Burkina Faso."),
    "",
    ("Un prototype web complet a été conçu et livré, couvrant l'intégralité des interfaces "
     "clients, agents et administrateurs, avec intégration Firebase (authentification anonyme, "
     "authentification email/mot de passe pour les agents, base de données Firestore en temps réel). "
     "Le présent devis couvre l'ensemble du travail réalisé et les modules en cours de finalisation."),
    "",
    "Stack technique : React 18 + Vite · Firebase Auth & Firestore · Tailwind CSS · "
    "React Router v6 · Recharts · PWA (manifest + service worker) · Déploiement Vercel.",
]
for i, line in enumerate(lines_ob):
    p4 = c.paragraphs[0] if i == 0 else c.add_paragraph()
    p4.paragraph_format.space_before = Pt(2)
    p4.paragraph_format.space_after  = Pt(2)
    run4 = p4.add_run(line)
    run4.font.size  = Pt(10)
    run4.font.name  = "Calibri"
    run4.font.color.rgb = DARK

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — DÉVELOPPEMENT
# ═══════════════════════════════════════════════════════════════════════════════
section_title("1. Développement — Prix fixe par module")
small_text(
    "Le développement est structuré en 5 modules livrables. "
    "Chaque module a fait l'objet d'une livraison et d'une validation.",
    italic=True
)

phases = [
    ("1",
     [("Prototype web — Interface client complète  ✓ Livré", True, DARK),
      ("Page d'accueil (pavé numérique, sélection plateforme), flux dépôt 6 étapes avec codes USSD "
       "(Orange Money / Telmob / Telecel), flux retrait 6 étapes avec calcul des frais, "
       "historique des transactions (Firebase Anonymous Auth).", False, GRAY)],
     "Interface client",
     "80 000 FCFA"),
    ("2",
     [("Espace Agent — Tableau de bord temps réel  ✓ Livré", True, DARK),
      ("Dashboard agent avec commandes en temps réel (onSnapshot), onglets Attente / En cours / Tout, "
       "alerte violette « Paiement envoyé », actions Accepter / Terminer / Annuler, "
       "historique des commandes traitées (/agent/orders), page login agent (numéro + mot de passe).", False, GRAY)],
     "Espace Agent",
     "55 000 FCFA"),
    ("3",
     [("Espace Admin — Dashboard analytique  ✓ Livré", True, DARK),
      ("6 KPIs (total, volumes dépôts/retraits, agents actifs, volume total), graphique en barres 7 jours, "
       "liste transactions complète avec filtres (statut, type, recherche), "
       "pages Dépôts et Retraits séparées, gestion des agents (création compte avec nom + numéro + mot de passe, "
       "activation/désactivation), sidebar structurée (Général / Activité / Équipe).", False, GRAY)],
     "Espace Admin",
     "65 000 FCFA"),
    ("4",
     [("Authentification & Sécurité  ✓ Livré", True, DARK),
      ("Firebase Anonymous Auth (clients sans inscription), Firebase Email/Password (agents), "
       "instance Firebase secondaire pour créer des agents sans déconnecter l'admin, "
       "AuthContext global avec signIn / logout / createAgentAccount / loginAgent, "
       "enregistrement du profil utilisateur dans Firestore.", False, GRAY)],
     "Auth & Sécurité",
     "35 000 FCFA"),
    ("5",
     [("PWA, déploiement & optimisations  ✓ Livré", True, DARK),
      ("Manifest PWA (manifest.json), service worker (cache offline), icônes 192×512 px, "
       "bannière d'installation personnalisée (beforeinstallprompt), "
       "déploiement Vercel avec vercel.json (SPA routing), favicon SVG.", False, GRAY)],
     "PWA & Déploiement",
     "20 000 FCFA"),
]

t1 = doc.add_table(rows=1 + len(phases) + 1, cols=4)
t1.style = "Table Grid"

# En-tête
for i, h in enumerate(["Ph.", "Contenu du module", "Titre", "Montant"]):
    shade_cell(t1.rows[0].cells[i], "1e3a8a")
    cell_text(t1.rows[0].cells[i], h, bold=True, color=WHITE, size=10,
              align=WD_ALIGN_PARAGRAPH.RIGHT if h == "Montant" else WD_ALIGN_PARAGRAPH.LEFT)

# Corps
for ri, (num, content_lines, titre, montant) in enumerate(phases, 1):
    if ri % 2 == 0:
        for ci in range(4):
            shade_cell(t1.rows[ri].cells[ci], "f9fafb")
    cell_text(t1.rows[ri].cells[0], num, size=12, bold=True, color=GRAY)
    # Cellule contenu multi-lignes
    c2 = t1.rows[ri].cells[1]
    c2.text = ""
    for j, (txt, bold, col) in enumerate(content_lines):
        p5 = c2.paragraphs[0] if j == 0 else c2.add_paragraph()
        p5.paragraph_format.space_before = Pt(2)
        p5.paragraph_format.space_after  = Pt(2)
        r5 = p5.add_run(txt)
        r5.font.bold  = bold
        r5.font.size  = Pt(10)
        r5.font.name  = "Calibri"
        r5.font.color.rgb = col
    cell_text(t1.rows[ri].cells[2], titre, size=10)
    cell_text(t1.rows[ri].cells[3], montant, bold=True, color=BLUE,
              align=WD_ALIGN_PARAGRAPH.RIGHT, size=11)

# Footer
shade_cell(t1.rows[-1].cells[0], "1e3a8a")
shade_cell(t1.rows[-1].cells[1], "1e3a8a")
shade_cell(t1.rows[-1].cells[2], "1e3a8a")
shade_cell(t1.rows[-1].cells[3], "1e3a8a")
cell_text(t1.rows[-1].cells[0], "", color=WHITE)
cell_text(t1.rows[-1].cells[1], "Sous-total développement", bold=True, color=WHITE, size=10)
cell_text(t1.rows[-1].cells[2], "", color=WHITE)
cell_text(t1.rows[-1].cells[3], "255 000 FCFA", bold=True, color=WHITE,
          align=WD_ALIGN_PARAGRAPH.RIGHT, size=11)

# Largeurs
for ri2 in range(len(t1.rows)):
    t1.rows[ri2].cells[0].width = Cm(1.0)
    t1.rows[ri2].cells[1].width = Cm(10.0)
    t1.rows[ri2].cells[2].width = Cm(3.0)
    t1.rows[ri2].cells[3].width = Cm(2.5)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — INFRASTRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════
section_title("2. Infrastructure & Hébergement")
small_text(
    "Frais d'infrastructure pour le fonctionnement de la plateforme en production.",
    italic=True
)

add_table(
    ["Composant", "Détail", "Montant"],
    [
        ("Vercel (hébergement web)",
         "Déploiement SPA sur Vercel — plan Hobby gratuit. "
         "CDN mondial, HTTPS automatique, déploiements continus depuis Git.",
         "0 FCFA"),
        ("Firebase Spark Plan",
         "Plan gratuit Firebase — Auth anonyme + Email/Password, Firestore (1 Go), "
         "Hosting. Suffisant pour la phase de lancement.",
         "0 FCFA"),
        ("Nom de domaine personnalisé",
         "Domaine .com ou .bf (ex : easypay.bf) — optionnel, "
         "à renouveler annuellement (≈ 5 000 – 15 000 FCFA/an selon registrar).",
         "Sur demande"),
        ("Firebase Blaze (si SMS OTP activé)",
         "Si l'authentification client par SMS est ajoutée, passage au plan payant Firebase. "
         "Coût variable selon volume SMS (≈ 0,006 $/SMS).",
         "Variable"),
    ],
    col_widths=[4.0, 10.0, 2.5],
    footer_rows=[("", "Sous-total infrastructure (lancement)", "0 FCFA")]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — OPTIONS
# ═══════════════════════════════════════════════════════════════════════════════
section_title("3. Options et évolutions futures")
small_text(
    "Ces modules s'ajoutent au forfait de base sur demande. Chaque option fait l'objet d'un avenant.",
    italic=True
)

add_table(
    ["Option", "Description", "Supplément"],
    [
        ("Authentification client par SMS (OTP)",
         "Remplacement de l'auth anonyme par une connexion via numéro de téléphone + SMS OTP. "
         "Permet au client de retrouver son historique sur plusieurs appareils.",
         "25 000 FCFA"),
        ("Notifications push Web (FCM)",
         "Alertes push navigateur : notification agent pour nouvelle commande, "
         "notification client pour changement de statut (accepté, terminé, annulé).",
         "20 000 FCFA"),
        ("Export CSV / PDF (Admin)",
         "Téléchargement des transactions filtrées en CSV ou en PDF mis en page, "
         "depuis le tableau de bord administrateur.",
         "15 000 FCFA"),
        ("Programme de parrainage",
         "Système de codes de parrainage : le client obtient un code unique, "
         "les filleuls sont tracés, tableau de bord parrainage dans l'espace client.",
         "30 000 FCFA"),
        ("Firestore Security Rules + Route Guards",
         "Mise en place des règles de sécurité Firestore par rôle (client / agent / admin) "
         "et protection des routes /agent et /admin par guards d'authentification.",
         "10 000 FCFA"),
        ("Maintenance évolutive (mensuelle)",
         "Support technique, corrections de bugs, mises à jour Firebase/dépendances, "
         "ajout de petites fonctionnalités (< 4h/mois).",
         "15 000 FCFA/mois"),
    ],
    col_widths=[4.5, 9.5, 2.5],
    footer_rows=[("", "Maximum toutes options activées (hors maintenance)", "100 000 FCFA")]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — RÉCAPITULATIF
# ═══════════════════════════════════════════════════════════════════════════════
section_title("4. Récapitulatif financier")

rt = doc.add_table(rows=4, cols=2)
rt.style = "Table Grid"

summary = [
    ("Développement — 5 modules livrés",   "255 000 FCFA"),
    ("Infrastructure — lancement",          "0 FCFA"),
    ("Forfait de base (hors options)",      "255 000 FCFA"),
    ("Options si toutes activées",          "+ 100 000 FCFA"),
]
colors = [DARK, DARK, BLUE_DARK, GRAY]
for ri, ((label, val), col) in enumerate(zip(summary, colors)):
    is_total = ri == 2
    bg = "dbeafe" if is_total else "ffffff"
    shade_cell(rt.rows[ri].cells[0], bg)
    shade_cell(rt.rows[ri].cells[1], bg)
    cell_text(rt.rows[ri].cells[0], label, bold=is_total, color=col, size=11)
    cell_text(rt.rows[ri].cells[1], val,   bold=is_total, color=col, size=12,
              align=WD_ALIGN_PARAGRAPH.RIGHT)
    rt.rows[ri].cells[0].width = Cm(12)
    rt.rows[ri].cells[1].width = Cm(4.5)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# Encadré total
tot = doc.add_table(rows=1, cols=1)
tot.style = "Table Grid"
shade_cell(tot.rows[0].cells[0], "1e3a8a")
c_tot = tot.rows[0].cells[0]
c_tot.text = ""
p_tot = c_tot.paragraphs[0]
p_tot.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tot.paragraph_format.space_before = Pt(6)
p_tot.paragraph_format.space_after  = Pt(2)
r_tot = p_tot.add_run("TOTAL MAXIMUM (toutes options)    355 000 FCFA")
r_tot.font.bold  = True
r_tot.font.size  = Pt(14)
r_tot.font.name  = "Calibri"
r_tot.font.color.rgb = WHITE

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — PAIEMENT
# ═══════════════════════════════════════════════════════════════════════════════
section_title("5. Conditions de paiement")

add_table(
    ["Échéance", "Condition de déclenchement", "Montant"],
    [
        ("Acompte — 40 %",
         "À la signature du présent devis. Couvre les modules 1, 2 et 3 (déjà livrés).",
         "102 000 FCFA"),
        ("Tranche 2 — 35 %",
         "Livraison et validation du module d'authentification (module 4) "
         "et des tests de bout en bout.",
         "89 250 FCFA"),
        ("Solde — 25 %",
         "Livraison finale, recette complète, déploiement production validé "
         "et remise des accès (Firebase, Vercel, code source).",
         "63 750 FCFA"),
    ],
    col_widths=[3.5, 10.0, 3.0],
    footer_rows=[("Total forfait de base", "", "255 000 FCFA")]
)

# Note jaune
note_t = doc.add_table(rows=1, cols=1)
note_t.style = "Table Grid"
shade_cell(note_t.rows[0].cells[0], YELLOW_BG)
nc = note_t.rows[0].cells[0]
nc.text = ""
np_ = nc.paragraphs[0]
np_.paragraph_format.space_before = Pt(4)
np_.paragraph_format.space_after  = Pt(4)
nr = np_.add_run(
    "Note : Les montants ci-dessus sont calculés sur le forfait de base de 255 000 FCFA. "
    "Tout module optionnel activé fera l'objet d'un avenant au contrat. "
    "Paiement accepté par Orange Money, Moov Money ou virement bancaire."
)
nr.font.size  = Pt(10)
nr.font.name  = "Calibri"
nr.font.color.rgb = RGBColor(0x78, 0x35, 0x0f)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — CONDITIONS
# ═══════════════════════════════════════════════════════════════════════════════
section_title("6. Conditions générales")

cg = doc.add_table(rows=1, cols=2)
cg.style = "Table Grid"
inc_cell = cg.rows[0].cells[0]
exc_cell = cg.rows[0].cells[1]
shade_cell(inc_cell, GREEN_BG)
shade_cell(exc_cell, "fff5f5")
inc_cell.width = Cm(8.25)
exc_cell.width = Cm(8.25)

def fill_cond(cell, title, items, col):
    cell.text = ""
    p6 = cell.paragraphs[0]
    p6.paragraph_format.space_before = Pt(4)
    r6 = p6.add_run(title)
    r6.font.bold  = True
    r6.font.size  = Pt(11)
    r6.font.name  = "Calibri"
    r6.font.color.rgb = col
    for item in items:
        pi = cell.add_paragraph()
        pi.paragraph_format.space_before = Pt(1)
        pi.paragraph_format.space_after  = Pt(1)
        ri = pi.add_run(item)
        ri.font.size  = Pt(10)
        ri.font.name  = "Calibri"
        ri.font.color.rgb = DARK

fill_cond(inc_cell, "✓  Inclus dans le devis", [
    "  Tous les modules livrés (1 à 5)",
    "  Code source complet versionné sur Git",
    "  Déploiement Vercel + configuration Firebase",
    "  PWA installable (Android, Desktop)",
    "  Icônes 192 px et 512 px générées",
    "  Documentation technique (ce cahier des charges)",
    "  Corrections de bugs pendant 60 jours après livraison finale",
    "  Formation à l'utilisation (1/2 journée distanciel)",
], GREEN)

fill_cond(exc_cell, "✗  Non inclus", [
    "  Frais SMS Firebase (si OTP activé)",
    "  Nom de domaine et renouvellement annuel",
    "  Firebase Blaze (si volume SMS dépasse le quota gratuit)",
    "  Maintenance évolutive après 60 jours (devis séparé)",
    "  Fonctionnalités hors périmètre actuel",
    "  Refonte de la charte graphique",
    "  Formation de nouveaux collaborateurs",
    "  Intégration API tierces non mentionnées",
], RED)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — SIGNATURE
# ═══════════════════════════════════════════════════════════════════════════════
section_title("7. Acceptation du devis")

p7 = doc.add_paragraph()
p7.paragraph_format.space_after = Pt(10)
r7 = p7.add_run(
    "Bon pour accord — En signant ce document, le client accepte les conditions du présent devis "
    "et s'engage à verser l'acompte de 40 % (102 000 FCFA) pour démarrer les travaux."
)
r7.font.size  = Pt(10)
r7.font.name  = "Calibri"
r7.font.color.rgb = DARK

# Tableau signature
sig = doc.add_table(rows=3, cols=2)
sig.style = "Table Grid"

# En-têtes
for i, h in enumerate(["Le prestataire", "Le client"]):
    shade_cell(sig.rows[0].cells[i], "1e3a8a")
    cell_text(sig.rows[0].cells[i], h, bold=True, color=WHITE, size=11,
              align=WD_ALIGN_PARAGRAPH.CENTER)

# Infos
cell_multiline(sig.rows[1].cells[0], [
    ("WOUBI ABDOUL AZIZ",            True,  DARK),
    ("Développeur Full-Stack & Mobile", False, GRAY),
    (f"Date : {TODAY_SHORT}",        False, DARK),
], size=11)
cell_multiline(sig.rows[1].cells[1], [
    ("[NOM DU CLIENT]",              True,  DARK),
    ("",                             False, DARK),
    ("Date : _____ / _____ / _____", False, DARK),
], size=11)

# Ligne de signature
for i in range(2):
    shade_cell(sig.rows[2].cells[i], "f9fafb")
    cell_text(sig.rows[2].cells[i], "Signature : _______________________________",
              size=10, color=GRAY)
    sig.rows[2].cells[i].height = Cm(2.5)

doc.add_paragraph()

# ── PIED DE PAGE ──────────────────────────────────────────────────────────────
ft = doc.add_paragraph()
ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
ft.paragraph_format.space_before = Pt(10)
pPr2 = ft._p.get_or_add_pPr()
pBdr2 = OxmlElement("w:pBdr")
top2 = OxmlElement("w:top")
top2.set(qn("w:val"),   "single")
top2.set(qn("w:sz"),    "4")
top2.set(qn("w:space"), "1")
top2.set(qn("w:color"), "e5e7eb")
pBdr2.append(top2)
pPr2.append(pBdr2)
rf = ft.add_run(
    f"Devis N° WAA-EASYPAY-2026-001  •  Valable 30 jours à compter du {TODAY}  •  "
    "WOUBI ABDOUL AZIZ  •  woubiaziz@gmail.com  •  +226 54 12 56 37"
)
rf.font.size   = Pt(8)
rf.font.name   = "Calibri"
rf.font.italic = True
rf.font.color.rgb = GRAY

# ── Sauvegarde ────────────────────────────────────────────────────────────────
out = r"C:\Users\Z book\Desktop\Devis_WAA-EASYPAY-2026-001.docx"
doc.save(out)
print(f"Devis genere : {out}")
