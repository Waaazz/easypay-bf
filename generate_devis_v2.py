from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

REF = "WAA-EASYPAY-2026-002"

# ── Marges ────────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Palette ───────────────────────────────────────────────────────────────────
BLUE      = RGBColor(0x25, 0x63, 0xeb)
BLUE_DARK = RGBColor(0x1e, 0x3a, 0x8a)
GREEN     = RGBColor(0x16, 0xa3, 0x4a)
RED       = RGBColor(0xdc, 0x26, 0x26)
GRAY      = RGBColor(0x6b, 0x72, 0x80)
DARK      = RGBColor(0x1a, 0x1a, 0x1a)
WHITE     = RGBColor(0xff, 0xff, 0xff)
YELLOW_BG = "fffbeb"
GREEN_BG  = "f0fdf4"
BLUE_BG   = "f8faff"

TODAY = datetime.date.today().strftime("%d %B %Y")
TODAY_SHORT = datetime.date.today().strftime("%d/%m/%Y")

# ── Helpers ───────────────────────────────────────────────────────────────────
def shade_cell(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, color=DARK, size=10,
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.size   = Pt(size)
    run.font.name   = "Calibri"
    run.font.color.rgb = color

def cell_multiline(cell, lines, size=10):
    cell.text = ""
    for i, (text, bold, color) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(text)
        run.font.bold  = bold
        run.font.size  = Pt(size)
        run.font.name  = "Calibri"
        run.font.color.rgb = color

def section_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.bold  = True
    run.font.size  = Pt(12)
    run.font.name  = "Calibri"
    run.font.color.rgb = BLUE
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "e5e7eb")
    pBdr.append(bot)
    pPr.append(pBdr)

def small_text(text, color=GRAY, size=10, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.name   = "Calibri"
    run.font.italic = italic
    run.font.color.rgb = color

def add_table(headers, rows, col_widths=None,
              header_bg="1e3a8a", alt_bg="f9fafb",
              footer_rows=None):
    t = doc.add_table(rows=1 + len(rows) + (len(footer_rows) if footer_rows else 0),
                      cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        shade_cell(t.rows[0].cells[i], header_bg)
        cell_text(t.rows[0].cells[i], h, bold=True, color=WHITE, size=10)
    for ri, row in enumerate(rows, 1):
        if ri % 2 == 0:
            for ci in range(len(headers)):
                shade_cell(t.rows[ri].cells[ci], alt_bg)
        for ci, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.RIGHT if ci == len(row)-1 else WD_ALIGN_PARAGRAPH.LEFT
            bold = ci == len(row)-1
            color = BLUE if bold else DARK
            cell_text(t.rows[ri].cells[ci], str(val), bold=bold, color=color,
                      align=align, size=10)
    if footer_rows:
        base = 1 + len(rows)
        for fi, fr in enumerate(footer_rows):
            for ci in range(len(headers)):
                shade_cell(t.rows[base+fi].cells[ci], "1e3a8a")
            for ci, val in enumerate(fr):
                align = WD_ALIGN_PARAGRAPH.RIGHT if ci == len(fr)-1 else WD_ALIGN_PARAGRAPH.LEFT
                cell_text(t.rows[base+fi].cells[ci], val, bold=True,
                          color=WHITE, size=10, align=align)
    if col_widths:
        for ri2 in range(len(t.rows)):
            for ci2, w in enumerate(col_widths):
                t.rows[ri2].cells[ci2].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t

# ═══════════════════════════════════════════════════════════════════════════════
#  EN-TÊTE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
run = p.add_run(
    f"DEVIS  |  Application Web ApollonPay / EasyPay BF  |  "
    f"Réf. {REF}  |  {TODAY}  |  Document confidentiel"
)
run.font.size   = Pt(8)
run.font.name   = "Calibri"
run.font.italic = True
run.font.color.rgb = GRAY

ht = doc.add_table(rows=1, cols=2)
ht.style = "Table Grid"
left  = ht.rows[0].cells[0]
right = ht.rows[0].cells[1]

left.paragraphs[0].paragraph_format.space_before = Pt(4)
r = left.paragraphs[0].add_run("WOUBI ABDOUL AZIZ")
r.font.bold = True; r.font.size = Pt(14); r.font.name = "Calibri"; r.font.color.rgb = DARK

for line_txt, b, sz, col in [
    ("Développeur Full-Stack & Mobile", False, 10, GRAY),
    ("", False, 6, DARK),
    ("woubiaziz@gmail.com", False, 10, DARK),
    ("+226 54 12 56 37 / +226 68 87 71 82", False, 10, DARK),
    ("Ouagadougou, Burkina Faso", False, 10, DARK),
    ("Licence en Génie Logiciel — UVBF", False, 10, DARK),
]:
    p2 = left.add_paragraph(line_txt)
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after  = Pt(1)
    for run2 in p2.runs:
        run2.font.bold  = b
        run2.font.size  = Pt(sz)
        run2.font.name  = "Calibri"
        run2.font.color.rgb = col

rp = right.paragraphs[0]
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r2 = rp.add_run("DEVIS")
r2.font.bold = True; r2.font.size = Pt(32); r2.font.name = "Calibri"; r2.font.color.rgb = DARK

for line_txt, b, sz, col in [
    (f"Réf. : {REF}", True, 10, BLUE),
    (f"Date : {TODAY}", False, 10, GRAY),
    ("Validité : 30 jours", False, 10, GRAY),
    ("", False, 6, DARK),
    ("CLIENT", True, 8, GRAY),
    ("[NOM DU CLIENT]", True, 12, DARK),
    ("Ouagadougou, Burkina Faso", False, 10, GRAY),
]:
    p3 = right.add_paragraph(line_txt)
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.paragraph_format.space_before = Pt(1)
    p3.paragraph_format.space_after  = Pt(1)
    for run3 in p3.runs:
        run3.font.bold  = b
        run3.font.size  = Pt(sz)
        run3.font.name  = "Calibri"
        run3.font.color.rgb = col

ht.rows[0].cells[0].width = Cm(9)
ht.rows[0].cells[1].width = Cm(7)
shade_cell(left, "ffffff")
shade_cell(right, "f8faff")

sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(8)
sep.paragraph_format.space_after  = Pt(8)
pPr = sep._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot = OxmlElement("w:bottom")
bot.set(qn("w:val"),   "single")
bot.set(qn("w:sz"),    "12")
bot.set(qn("w:space"), "1")
bot.set(qn("w:color"), "2563eb")
pBdr.append(bot)
pPr.append(pBdr)

# ── OBJET ─────────────────────────────────────────────────────────────────────
section_title("Objet de la prestation")

ob = doc.add_table(rows=1, cols=1)
ob.style = "Table Grid"
shade_cell(ob.rows[0].cells[0], BLUE_BG)
c = ob.rows[0].cells[0]

lines_ob = [
    "Développement d'une application web mobile-first (PWA) — ApollonPay / EasyPay BF — "
    "plateforme de dépôt, retrait et recharge d'abonnement dédiée aux paris sportifs "
    "(1XBET, MelBet, Betwinner) et aux abonnements télé (CANAL+, Canalbox) au Burkina Faso.",
    "",
    "Le présent devis couvre l'intégralité de l'application livrée à ce jour : interfaces "
    "client, agent et administrateur, moteur de transactions temps réel avec Firebase/"
    "Firestore, assignation intelligente des agents, module d'abonnement CANAL+/Canalbox, "
    "et intégration PWA installable directement depuis le navigateur (sans passer par un "
    "store d'applications).",
    "",
    "Stack technique : React 18 + Vite · Firebase Auth & Firestore (temps réel, transactions "
    "atomiques) · Tailwind CSS · React Router v6 · Recharts · PWA (manifest + service worker "
    "+ invite d'installation) · Déploiement Vercel.",
]
for i, line in enumerate(lines_ob):
    p4 = c.paragraphs[0] if i == 0 else c.add_paragraph()
    p4.paragraph_format.space_before = Pt(2)
    p4.paragraph_format.space_after  = Pt(2)
    run4 = p4.add_run(line)
    run4.font.size  = Pt(10)
    run4.font.name  = "Calibri"
    run4.font.color.rgb = DARK

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — DÉVELOPPEMENT (forfait fixe global)
# ═══════════════════════════════════════════════════════════════════════════════
section_title("1. Développement — Forfait fixe global : 750 000 FCFA")
small_text(
    "Prix fixe et global pour l'ensemble de l'application, réparti ci-dessous par module "
    "à titre indicatif. Ce forfait n'évolue pas en fonction du temps réellement passé.",
    italic=True
)

modules = [
    ("1",
     [("Interface Client", True, DARK),
      ("Page d'accueil (pavé numérique, sélection plateforme), flux dépôt et retrait "
       "multi-étapes avec calcul des frais et codes USSD par opérateur, historique des "
       "transactions, authentification et inscription client.", False, GRAY)],
     "150 000 FCFA"),
    ("2",
     [("Module Abonnements CANAL+ / Canalbox", True, DARK),
      ("Parcours complet dédié : saisie du décodeur, écran de vérification du compte "
       "(informations personnelles + abonnement en cours), sélection de l'offre, de la "
       "durée et des options, choix du moyen de paiement, feuille de paiement et "
       "confirmation.", False, GRAY)],
     "120 000 FCFA"),
    ("3",
     [("Espace Agent", True, DARK),
      ("Tableau de bord temps réel (commandes en attente / en cours / historique), bascule "
       "de disponibilité, traitement des commandes en un clic, alertes et notifications "
       "navigateur pour les nouvelles demandes.", False, GRAY)],
     "130 000 FCFA"),
    ("4",
     [("Espace Admin", True, DARK),
      ("Tableau de bord analytique (KPIs, graphique d'activité), gestion complète des "
       "transactions (filtres, recherche), gestion des comptes agents (création, "
       "activation/désactivation, numéros par opérateur).", False, GRAY)],
     "130 000 FCFA"),
    ("5",
     [("Moteur transactionnel & logique métier", True, DARK),
      ("Transactions Firestore atomiques (anti double-traitement), assignation "
       "intelligente des agents par opérateur avec équilibrage de charge, gestion de la "
       "disponibilité en temps réel, repli en diffusion contrôlée en l'absence d'agent "
       "disponible.", False, GRAY)],
     "120 000 FCFA"),
    ("6",
     [("Authentification & Sécurité", True, DARK),
      ("Firebase Auth (accès client, comptes agents et admin), gestion des profils "
       "utilisateurs et des rôles, routes protégées par rôle.", False, GRAY)],
     "60 000 FCFA"),
    ("7",
     [("PWA, déploiement & optimisations", True, DARK),
      ("Manifest PWA, service worker, icônes, bannière d'installation personnalisée "
       "(écran d'accueil, sans store), déploiement Vercel.", False, GRAY)],
     "40 000 FCFA"),
]

t1 = doc.add_table(rows=1 + len(modules) + 1, cols=3)
t1.style = "Table Grid"
for i, h in enumerate(["Contenu du module", "Titre", "Montant"]):
    shade_cell(t1.rows[0].cells[i], "1e3a8a")
    cell_text(t1.rows[0].cells[i], h, bold=True, color=WHITE, size=10)

for ri, (num, content_lines, montant) in enumerate(modules, 1):
    if ri % 2 == 0:
        for ci in range(3):
            shade_cell(t1.rows[ri].cells[ci], "f9fafb")
    c2 = t1.rows[ri].cells[0]
    c2.text = ""
    for j, (txt, bold, col) in enumerate(content_lines):
        p5 = c2.paragraphs[0] if j == 0 else c2.add_paragraph()
        p5.paragraph_format.space_before = Pt(2)
        p5.paragraph_format.space_after  = Pt(2)
        r5 = p5.add_run(txt)
        r5.font.bold  = bold
        r5.font.size  = Pt(10)
        r5.font.name  = "Calibri"
        r5.font.color.rgb = col
    cell_text(t1.rows[ri].cells[1], f"Module {num}", size=10)
    cell_text(t1.rows[ri].cells[2], montant, bold=True, color=BLUE,
              align=WD_ALIGN_PARAGRAPH.RIGHT, size=11)

shade_cell(t1.rows[-1].cells[0], "1e3a8a")
shade_cell(t1.rows[-1].cells[1], "1e3a8a")
shade_cell(t1.rows[-1].cells[2], "1e3a8a")
cell_text(t1.rows[-1].cells[0], "Total forfait — Développement", bold=True, color=WHITE, size=10)
cell_text(t1.rows[-1].cells[1], "", color=WHITE)
cell_text(t1.rows[-1].cells[2], "750 000 FCFA", bold=True, color=WHITE,
          align=WD_ALIGN_PARAGRAPH.RIGHT, size=11)

for ri2 in range(len(t1.rows)):
    t1.rows[ri2].cells[0].width = Cm(11.0)
    t1.rows[ri2].cells[1].width = Cm(2.5)
    t1.rows[ri2].cells[2].width = Cm(2.5)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — INFRASTRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════
section_title("2. Infrastructure & Hébergement")
small_text(
    "Frais de fonctionnement de la plateforme en production — non compris dans le forfait "
    "de développement, à la charge du client.",
    italic=True
)

add_table(
    ["Composant", "Détail", "Montant"],
    [
        ("Vercel (hébergement web)",
         "Déploiement sur Vercel — plan Hobby gratuit. CDN mondial, HTTPS automatique.",
         "0 FCFA"),
        ("Firebase Spark Plan",
         "Plan gratuit Firebase — Auth, Firestore, Hosting. Suffisant au lancement.",
         "0 FCFA"),
        ("Nom de domaine personnalisé",
         "Optionnel — à renouveler annuellement (≈ 5 000 – 15 000 FCFA/an selon registrar).",
         "Sur demande"),
        ("Firebase Blaze (si volume élevé)",
         "Passage au plan payant si les quotas gratuits Firestore/Auth sont dépassés.",
         "Variable"),
    ],
    col_widths=[4.0, 10.0, 2.5],
    footer_rows=[("", "Sous-total infrastructure (lancement)", "0 FCFA")]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — ÉVOLUTIONS FUTURES (hors forfait)
# ═══════════════════════════════════════════════════════════════════════════════
section_title("3. Évolutions futures — hors forfait, sur devis séparé")
small_text(
    "Non comprises dans le forfait de 750 000 FCFA. Ces travaux dépendent de partenariats "
    "externes (InTouch, CORIS Money, LigdiCash, MobCash...) et feront l'objet d'un devis "
    "distinct une fois les accords commerciaux obtenus.",
    italic=True
)

add_table(
    ["Évolution", "Description"],
    [
        ("Vérification décodeur CANAL+/Canalbox en temps réel",
         "Intégration API partenaire (InTouch, CORIS Money...) pour remplacer la saisie "
         "manuelle par une vérification automatique du compte abonné."),
        ("Paiement automatisé bookmakers (1XBET, Betwinner, MelBet)",
         "Intégration API partenaire (LigdiCash, MobCash Super Agent...) pour la "
         "vérification d'ID joueur et le crédit automatique du compte de pari."),
        ("Maintenance évolutive",
         "Support, corrections et ajout de fonctionnalités après la période de garantie — "
         "forfait mensuel à définir."),
    ],
    col_widths=[6.0, 10.5],
)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — RÉCAPITULATIF
# ═══════════════════════════════════════════════════════════════════════════════
section_title("4. Récapitulatif financier")

tot = doc.add_table(rows=1, cols=1)
tot.style = "Table Grid"
shade_cell(tot.rows[0].cells[0], "1e3a8a")
c_tot = tot.rows[0].cells[0]
c_tot.text = ""
p_tot = c_tot.paragraphs[0]
p_tot.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tot.paragraph_format.space_before = Pt(6)
p_tot.paragraph_format.space_after  = Pt(2)
r_tot = p_tot.add_run("FORFAIT GLOBAL FIXE    750 000 FCFA")
r_tot.font.bold  = True
r_tot.font.size  = Pt(16)
r_tot.font.name  = "Calibri"
r_tot.font.color.rgb = WHITE

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — PAIEMENT
# ═══════════════════════════════════════════════════════════════════════════════
section_title("5. Conditions de paiement")

add_table(
    ["Échéance", "Condition de déclenchement", "Montant"],
    [
        ("Acompte — 30 %",
         "À la signature du présent devis.",
         "225 000 FCFA"),
        ("Tranche 2 — 35 %",
         "Validation des modules Client, Agent et Admin par le client.",
         "262 500 FCFA"),
        ("Solde — 35 %",
         "Livraison finale, recette complète et remise des accès "
         "(Firebase, Vercel, code source).",
         "262 500 FCFA"),
    ],
    col_widths=[3.5, 10.0, 3.0],
    footer_rows=[("Total forfait", "", "750 000 FCFA")]
)

note_t = doc.add_table(rows=1, cols=1)
note_t.style = "Table Grid"
shade_cell(note_t.rows[0].cells[0], YELLOW_BG)
nc = note_t.rows[0].cells[0]
nc.text = ""
np_ = nc.paragraphs[0]
np_.paragraph_format.space_before = Pt(4)
np_.paragraph_format.space_after  = Pt(4)
nr = np_.add_run(
    "Note : 750 000 FCFA est un prix fixe et global, non révisable en cours de projet sauf "
    "ajout de fonctionnalités hors périmètre (voir section 3). Paiement accepté par Orange "
    "Money, Moov Money ou virement bancaire."
)
nr.font.size  = Pt(10)
nr.font.name  = "Calibri"
nr.font.color.rgb = RGBColor(0x78, 0x35, 0x0f)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — CONDITIONS
# ═══════════════════════════════════════════════════════════════════════════════
section_title("6. Conditions générales")

cg = doc.add_table(rows=1, cols=2)
cg.style = "Table Grid"
inc_cell = cg.rows[0].cells[0]
exc_cell = cg.rows[0].cells[1]
shade_cell(inc_cell, GREEN_BG)
shade_cell(exc_cell, "fff5f5")
inc_cell.width = Cm(8.25)
exc_cell.width = Cm(8.25)

def fill_cond(cell, title, items, col):
    cell.text = ""
    p6 = cell.paragraphs[0]
    p6.paragraph_format.space_before = Pt(4)
    r6 = p6.add_run(title)
    r6.font.bold  = True
    r6.font.size  = Pt(11)
    r6.font.name  = "Calibri"
    r6.font.color.rgb = col
    for item in items:
        pi = cell.add_paragraph()
        pi.paragraph_format.space_before = Pt(1)
        pi.paragraph_format.space_after  = Pt(1)
        ri = pi.add_run(item)
        ri.font.size  = Pt(10)
        ri.font.name  = "Calibri"
        ri.font.color.rgb = DARK

fill_cond(inc_cell, "Inclus dans le forfait", [
    "Les 7 modules décrits en section 1",
    "Code source complet versionné sur Git",
    "Déploiement Vercel + configuration Firebase",
    "PWA installable (Android, iOS Safari, Desktop) sans store",
    "Icônes et assets graphiques de l'application",
    "Corrections de bugs pendant 60 jours après livraison finale",
    "Formation à l'utilisation (1/2 journée à distance)",
], GREEN)

fill_cond(exc_cell, "Non inclus", [
    "Intégrations API partenaires (section 3)",
    "Frais SMS Firebase au-delà du quota gratuit",
    "Nom de domaine et renouvellement annuel",
    "Maintenance évolutive après 60 jours (devis séparé)",
    "Fonctionnalités hors périmètre du présent devis",
    "Refonte de la charte graphique / redesign complet",
    "Formation de nouveaux collaborateurs après livraison",
], RED)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — SIGNATURE
# ═══════════════════════════════════════════════════════════════════════════════
section_title("7. Acceptation du devis")

p7 = doc.add_paragraph()
p7.paragraph_format.space_after = Pt(10)
r7 = p7.add_run(
    "Bon pour accord — En signant ce document, le client accepte les conditions du présent "
    "devis et s'engage à verser l'acompte de 30 % (225 000 FCFA) pour démarrer les travaux."
)
r7.font.size  = Pt(10)
r7.font.name  = "Calibri"
r7.font.color.rgb = DARK

sig = doc.add_table(rows=3, cols=2)
sig.style = "Table Grid"
for i, h in enumerate(["Le prestataire", "Le client"]):
    shade_cell(sig.rows[0].cells[i], "1e3a8a")
    cell_text(sig.rows[0].cells[i], h, bold=True, color=WHITE, size=11,
              align=WD_ALIGN_PARAGRAPH.CENTER)

cell_multiline(sig.rows[1].cells[0], [
    ("WOUBI ABDOUL AZIZ",            True,  DARK),
    ("Développeur Full-Stack & Mobile", False, GRAY),
    (f"Date : {TODAY_SHORT}",        False, DARK),
], size=11)
cell_multiline(sig.rows[1].cells[1], [
    ("[NOM DU CLIENT]",              True,  DARK),
    ("",                             False, DARK),
    ("Date : _____ / _____ / _____", False, DARK),
], size=11)

for i in range(2):
    shade_cell(sig.rows[2].cells[i], "f9fafb")
    cell_text(sig.rows[2].cells[i], "Signature : _______________________________",
              size=10, color=GRAY)
    sig.rows[2].cells[i].height = Cm(2.5)

doc.add_paragraph()

# ── PIED DE PAGE ──────────────────────────────────────────────────────────────
ft = doc.add_paragraph()
ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
ft.paragraph_format.space_before = Pt(10)
pPr2 = ft._p.get_or_add_pPr()
pBdr2 = OxmlElement("w:pBdr")
top2 = OxmlElement("w:top")
top2.set(qn("w:val"),   "single")
top2.set(qn("w:sz"),    "4")
top2.set(qn("w:space"), "1")
top2.set(qn("w:color"), "e5e7eb")
pBdr2.append(top2)
pPr2.append(pBdr2)
rf = ft.add_run(
    f"Devis N° {REF}  •  Valable 30 jours à compter du {TODAY}  •  "
    "WOUBI ABDOUL AZIZ  •  woubiaziz@gmail.com  •  +226 54 12 56 37"
)
rf.font.size   = Pt(8)
rf.font.name   = "Calibri"
rf.font.italic = True
rf.font.color.rgb = GRAY

# ── Sauvegarde ────────────────────────────────────────────────────────────────
out = r"C:\Users\Z book\Desktop\Devis_WAA-EASYPAY-2026-002.docx"
doc.save(out)
print(f"Devis genere : {out}")
